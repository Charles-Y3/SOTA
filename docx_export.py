"""Save an edited transcript as .docx (when Microsoft Word is available on
the machine) or plain .txt otherwise."""

import os
import shutil
import sys

from transcriber import unique_path


def word_available():
    """True if Microsoft Word appears to be installed — Windows registry on
    Windows, the standard install locations on macOS (winreg doesn't exist
    there, so before this check Mac users with Word still always got .txt)."""
    if sys.platform == "darwin":
        return any(os.path.isdir(p) for p in (
            "/Applications/Microsoft Word.app",
            os.path.expanduser("~/Applications/Microsoft Word.app"),
        ))
    try:
        import winreg
    except ImportError:
        return False
    for root in (winreg.HKEY_CLASSES_ROOT, winreg.HKEY_LOCAL_MACHINE):
        subkey = "Word.Application" if root == winreg.HKEY_CLASSES_ROOT \
            else r"SOFTWARE\Microsoft\Office\Word"
        try:
            with winreg.OpenKey(root, subkey):
                return True
        except OSError:
            continue
    return False


def read_transcript(path):
    """Read a transcript back into editable text, handling .docx and .txt."""
    if path.lower().endswith(".docx"):
        from docx import Document

        doc = Document(path)
        return "\n\n".join(p.text for p in doc.paragraphs)
    with open(path, encoding="utf-8") as f:
        return f.read()


def _paragraphs(text):
    text = text.replace("\r\n", "\n").strip("\n")
    if not text.strip():
        return []
    # A blank line separates paragraphs; if there are none, treat each
    # non-empty line as its own paragraph (e.g. timestamped output).
    if "\n\n" in text:
        blocks = text.split("\n\n")
    else:
        blocks = text.split("\n")
    return [b.strip() for b in blocks if b.strip()]


def save_transcript(text, dest_folder, stem):
    """Write `text` into `dest_folder` as `stem.docx` or `stem.txt`.

    Returns (path, kind) where kind is "docx" or "txt".
    """
    os.makedirs(dest_folder, exist_ok=True)

    if word_available():
        try:
            from docx import Document

            doc = Document()
            paras = _paragraphs(text)
            if paras:
                for p in paras:
                    doc.add_paragraph(p)
            else:
                doc.add_paragraph("")
            path = unique_path(os.path.join(dest_folder, stem + ".docx"))
            doc.save(path)
            return path, "docx"
        except Exception:
            pass  # fall through to txt so the user never loses their edit

    path = unique_path(os.path.join(dest_folder, stem + ".txt"))
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)
    return path, "txt"


def edited_copy_path(dest_folder, stem):
    """Picks the path a live session's evolving '(edited)' copy will live
    at — same unique-name rule save_transcript uses (so it never collides
    with an unrelated existing file), computed once by the caller and then
    reused for every subsequent save_transcript_at call, instead of a fresh
    numbered name each time. Extension follows the same rule save_transcript
    uses (docx if Word is available, else txt), fixed for the rest of the
    session so a save never straddles two file kinds."""
    os.makedirs(dest_folder, exist_ok=True)
    ext = ".docx" if word_available() else ".txt"
    return unique_path(os.path.join(dest_folder, stem + ext))


def save_transcript_at(text, path):
    """Overwrites `path` with `text` in place, keeping exactly one backup
    of whatever was there before — `<name> (previous)<ext>`, itself
    overwritten on the next call rather than accumulating. One step of
    undo-safety for a live session's repeatedly-saved edited copy, without
    the multiple-file clutter a fresh numbered file per save would produce.
    Kind (docx vs txt) follows `path`'s own extension, not word_available()
    — set once by edited_copy_path and never re-decided mid-session."""
    if os.path.isfile(path):
        base, ext = os.path.splitext(path)
        try:
            shutil.copy2(path, f"{base} (previous){ext}")
        except Exception:
            pass  # best-effort; must never block saving the new content

    if path.lower().endswith(".docx"):
        from docx import Document

        doc = Document()
        paras = _paragraphs(text)
        if paras:
            for p in paras:
                doc.add_paragraph(p)
        else:
            doc.add_paragraph("")
        doc.save(path)
    else:
        with open(path, "w", encoding="utf-8") as f:
            f.write(text)


def append_transcript(path, paragraphs):
    """Appends `paragraphs` (a list of strings) to an existing transcript
    file written by save_transcript, preserving its one-block-per-paragraph
    structure — used by the Live tab's draft saves, where a session's
    transcript grows over multiple writes instead of being written once."""
    paras = [p.strip() for p in paragraphs if p and p.strip()]
    if not paras:
        return
    if path.lower().endswith(".docx"):
        from docx import Document

        doc = Document(path)
        for p in paras:
            doc.add_paragraph(p)
        doc.save(path)
        return
    # save_transcript leaves .txt files ending in a single newline, so each
    # appended block starts with one more to form the blank-line separator.
    with open(path, "a", encoding="utf-8") as f:
        for p in paras:
            f.write("\n" + p + "\n")
